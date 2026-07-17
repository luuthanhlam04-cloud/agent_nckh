"""
docx_exporter.py - Xuat Bao Cao Hoc Thuat ra Word (.docx) (Giai doan 5)
=========================================================================
Chuc nang chinh:
  - Nhan topic tu RegexInterceptor (intent=EXPORT_DOCX)
  - Goi RAG rong hon (top_k=10) de thu thap tri thuc
  - Dung python-docx xay dung bao cao Word co cau truc hoc thuat
  - Luu Desktop/{topic}_{datetime}.docx
  - Thong bao TTS: "Da ket xuat bao cao ra Desktop"

Risk 3 Fix (Export UX):
  - Tra ve (path, status_message) thay vi chi path
  - status_message duoc dung boi AIWorker de emit status_update signal
  - UI hien thi "Dang thu thap tri thuc va ket xuat MS Word..." trong luc cho

Kien truc Word document:
  - Heading 1: Ten bao cao
  - Metadata: Ngay tao, so nguon
  - Section 1: Tom tat dieu hanh (1 doan)
  - Section 2: Noi dung chinh (Markdown to Word paragraphs)
  - Section 3: Tai lieu tham khao (sources list)
  - Format: Font Times New Roman 12, gian dong 1.5, canh deu 2 ben
"""

import os
import gc
import logging
import re
from datetime import datetime
from pathlib import Path
from typing import Optional, Tuple

logger = logging.getLogger("DocxExporter")

# Thu muc luu file (Desktop cua nguoi dung)
_DESKTOP = os.path.join(os.path.expanduser("~"), "Desktop")

# Ky tu khong hop le trong ten file Windows
_INVALID_CHARS = re.compile(r'[\\/:*?"<>|]')


def _sanitize_filename(name: str) -> str:
    """Xoa ky tu khong hop le trong ten file."""
    cleaned = _INVALID_CHARS.sub("_", name)
    return cleaned[:60].strip()   # Gioi han do dai ten file


class DocxExporter:
    """
    Xuat bao cao hoc thuat ra file Word dua tren RAG query.

    Usage (tu process_user_input trong main.py):
        exporter = DocxExporter(orchestrator=orchestrator)
        path, msg = exporter.export(topic="GraphRAG")
    """

    def __init__(self, orchestrator):
        self._orchestrator = orchestrator
        logger.info("[DocxExporter] Khoi tao.")

    def export(self, topic: str) -> Tuple[str, str]:
        """
        Thuc hien toan bo pipeline export.

        Returns:
            (output_path, status_message)
            - output_path  : Duong dan file .docx da luu
            - status_message: Chuoi thong bao cho TTS va UI
        """
        if not topic or not topic.strip():
            return "", "Chu de bao cao khong duoc de trong."

        topic = topic.strip()
        logger.info("[DocxExporter] Bat dau xuat bao cao: '%s'", topic)

        # ── Buoc 1: Thu thap tri thuc qua RAG ───────────────────────────────
        logger.info("[DocxExporter] Dang truy van RAG (top_k=10)...")
        try:
            # Goi orchestrator voi query mo rong de lay nhieu context hon
            # [BUG-J FIX] orchestrator.run() trả về Generator, phải join hết chunks
            # Code cũ gán generator object trực tiếp → _build_docx() nhận iterator rỗng
            gen = self._orchestrator.run(
                user_input=(
                    f"Hay viet mot bai bao cao hoc thuat day du ve chu de: {topic}. "
                    f"Bao gom: dinh nghia, phuong phap, ung dung, han che va huong phat trien. "
                    f"Su dung tat ca tai lieu trong kho tri thuc."
                )
            )
            rag_content = "".join(chunk for chunk in gen if chunk)
            sources = self._orchestrator.get_last_sources()   # Lay nguon truy xuat
        except AttributeError:
            # Orchestrator khong co get_last_sources -> dung empty list
            gen = self._orchestrator.run(
                user_input=f"Hay viet mot bao cao hoc thuat ve: {topic}"
            )
            rag_content = "".join(chunk for chunk in gen if chunk)
            sources = []
        except Exception as e:
            logger.error("[DocxExporter] Loi RAG: %s", e, exc_info=True)
            return "", f"Loi truy xuat tri thuc: {str(e)[:80]}"

        # ── Buoc 2: Xay dung file Word ───────────────────────────────────────
        logger.info("[DocxExporter] Dang tao file Word...")
        output_path = self._build_docx(topic, rag_content, sources)

        gc.collect()   # Don RAM sau khi tao document object lon

        if not output_path:
            return "", "Loi tao file Word. Kiem tra python-docx da cai."

        # ── Buoc 3: Thong bao thanh cong ─────────────────────────────────────
        filename = Path(output_path).name
        status_msg = (
            f"Da ket xuat bao cao hoc thuat '{topic}' ra man hinh Desktop. "
            f"Ten file: {filename}"
        )
        logger.info("[DocxExporter] Thanh cong: %s", output_path)
        return output_path, status_msg

    def _build_docx(
        self,
        topic: str,
        content: str,
        sources: list,
    ) -> Optional[str]:
        """
        Xay dung file .docx voi cau truc hoc thuat.
        Tra ve duong dan file hoac None neu that bai.
        """
        try:
            from docx import Document
            from docx.shared import Pt, Cm, RGBColor
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            from docx.oxml.ns import qn
            from docx.oxml import OxmlElement

            doc = Document()

            # ── Cau hinh trang (A4, le 2.5cm) ─────────────────────────────
            section = doc.sections[0]
            section.page_height  = Cm(29.7)
            section.page_width   = Cm(21.0)
            section.left_margin  = Cm(2.5)
            section.right_margin = Cm(2.5)
            section.top_margin   = Cm(2.5)
            section.bottom_margin = Cm(2.5)

            # ── Tieu de chinh (Heading 1) ─────────────────────────────────
            title_para = doc.add_paragraph()
            title_run = title_para.add_run(f"Báo Cáo Học Thuật: {topic}")
            title_run.font.name = "Times New Roman"
            title_run.font.size = Pt(16)
            title_run.font.bold = True
            title_run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x6C)  # Dark blue
            title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            self._set_east_asian_font(title_run, "Times New Roman")

            # ── Metadata ──────────────────────────────────────────────────
            meta_para = doc.add_paragraph()
            meta_text = (
                f"Ngày tạo: {datetime.now().strftime('%d/%m/%Y %H:%M')}  |  "
                f"Tạo bởi: Digital Scholar v3.0"
            )
            meta_run = meta_para.add_run(meta_text)
            meta_run.font.name  = "Times New Roman"
            meta_run.font.size  = Pt(10)
            meta_run.font.italic = True
            meta_run.font.color.rgb = RGBColor(0x80, 0x80, 0x80)
            meta_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            self._set_east_asian_font(meta_run, "Times New Roman")

            doc.add_paragraph()   # Khoang trang

            # ── Duong ke ngang ─────────────────────────────────────────────
            self._add_horizontal_rule(doc)

            # ── Noi dung chinh ─────────────────────────────────────────────
            doc.add_heading("Nội Dung Báo Cáo", level=1)
            self._add_content_paragraphs(doc, content)

            # ── Tai lieu tham khao ─────────────────────────────────────────
            if sources:
                doc.add_heading("Tài Liệu Tham Khảo", level=1)
                for i, src in enumerate(sources[:10], 1):
                    src_text = src if isinstance(src, str) else str(src)
                    ref_para = doc.add_paragraph(style="List Number")
                    ref_run = ref_para.add_run(src_text[:200])
                    ref_run.font.name = "Times New Roman"
                    ref_run.font.size = Pt(11)
                    self._set_east_asian_font(ref_run, "Times New Roman")

            # ── Footer ────────────────────────────────────────────────────
            doc.add_paragraph()
            footer_para = doc.add_paragraph()
            footer_run = footer_para.add_run(
                "--- Bao cao nay duoc tao tu dong boi Digital Scholar (Last Agent V3.0) ---"
            )
            footer_run.font.size = Pt(9)
            footer_run.font.italic = True
            footer_run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
            footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

            # ── Luu file ──────────────────────────────────────────────────
            safe_topic = _sanitize_filename(topic)
            timestamp  = datetime.now().strftime("%Y%m%d_%H%M")
            filename   = f"BaoCao_{safe_topic}_{timestamp}.docx"
            output_path = os.path.join(_DESKTOP, filename)

            Path(_DESKTOP).mkdir(parents=True, exist_ok=True)
            doc.save(output_path)
            return output_path

        except ImportError:
            logger.error(
                "[DocxExporter] python-docx chua cai. "
                "Chay: pip install python-docx"
            )
            return None
        except Exception as e:
            logger.error("[DocxExporter] Loi tao docx: %s", e, exc_info=True)
            return None

    def _add_content_paragraphs(self, doc, content: str):
        """
        Phan tich noi dung Markdown don gian va them vao doc.
        Ho tro: # Heading, ## Subheading, - bullet, plain text.
        """
        from docx.shared import Pt
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        lines = content.strip().split("\n")
        for line in lines:
            line = line.rstrip()
            if not line:
                doc.add_paragraph()
                continue

            # H1/H2/H3
            if line.startswith("### "):
                h = doc.add_heading(line[4:], level=3)
                self._format_heading(h, Pt(12))
            elif line.startswith("## "):
                h = doc.add_heading(line[3:], level=2)
                self._format_heading(h, Pt(13))
            elif line.startswith("# "):
                h = doc.add_heading(line[2:], level=1)
                self._format_heading(h, Pt(14))
            # Bullet
            elif line.startswith("- ") or line.startswith("* "):
                p = doc.add_paragraph(style="List Bullet")
                run = p.add_run(line[2:])
                run.font.name = "Times New Roman"
                run.font.size = Pt(12)
                self._set_east_asian_font(run, "Times New Roman")
            # Numbered list
            elif len(line) > 2 and line[0].isdigit() and line[1] in ".)" :
                p = doc.add_paragraph(style="List Number")
                run = p.add_run(line[2:].strip())
                run.font.name = "Times New Roman"
                run.font.size = Pt(12)
                self._set_east_asian_font(run, "Times New Roman")
            # Bold (**text**)
            elif "**" in line:
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                parts = line.split("**")
                for i, part in enumerate(parts):
                    if not part:
                        continue
                    r = p.add_run(part)
                    r.font.name = "Times New Roman"
                    r.font.size = Pt(12)
                    r.font.bold = (i % 2 == 1)   # Odd index = bold
                    self._set_east_asian_font(r, "Times New Roman")
            # Plain paragraph
            else:
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                p.paragraph_format.first_line_indent = Pt(24)
                p.paragraph_format.space_after = Pt(6)

                run = p.add_run(line)
                run.font.name = "Times New Roman"
                run.font.size = Pt(12)
                self._set_east_asian_font(run, "Times New Roman")

    def _format_heading(self, heading_para, font_size):
        """Ap dung font Times New Roman cho heading."""
        from docx.shared import Pt
        for run in heading_para.runs:
            run.font.name = "Times New Roman"
            run.font.size = font_size
            self._set_east_asian_font(run, "Times New Roman")

    @staticmethod
    def _set_east_asian_font(run, font_name: str):
        """
        Dat font cho ky tu Dong A (CJK + Vietnamese).
        python-docx thuong khong tu dong ap dung font cho ky tu Unicode.

        [M3-FIX] Import OxmlElement CUC BO ben trong ham thay vi dung bien
        module-level co the la None khi python-docx chua cai.
        Tat ca exception duoc bat im lang de khong anh huong chuc nang chinh.
        """
        try:
            from docx.oxml.ns import qn
            from docx.oxml import OxmlElement   # [M3-FIX] Import cuc bo, an toan
            rPr = run._r.get_or_add_rPr()
            rFonts = rPr.find(qn("w:rFonts"))
            if rFonts is None:
                rFonts = OxmlElement("w:rFonts")
                rPr.insert(0, rFonts)
            rFonts.set(qn("w:ascii"), font_name)
            rFonts.set(qn("w:hAnsi"), font_name)
            rFonts.set(qn("w:cs"), font_name)
            rFonts.set(qn("w:eastAsia"), font_name)
        except Exception as e:
            logger.debug("[DocxExporter] Khong the set east asian font: %s", e)

    @staticmethod
    def _add_horizontal_rule(doc):
        """Them duong ke ngang giua cac phan."""
        try:
            from docx.oxml.ns import qn
            from docx.oxml import OxmlElement

            p = doc.add_paragraph()
            pPr = p._p.get_or_add_pPr()
            pBdr = OxmlElement("w:pBdr")
            bottom = OxmlElement("w:bottom")
            bottom.set(qn("w:val"), "single")
            bottom.set(qn("w:sz"), "6")
            bottom.set(qn("w:space"), "1")
            bottom.set(qn("w:color"), "1A1A6C")
            pBdr.append(bottom)
            pPr.append(pBdr)
        except Exception as e:
            logger.debug("[DocxExporter] Khong the them duong ke: %s", e)


# [M3-FIX] OxmlElement module-level import da duoc xoa.
# _set_east_asian_font gio import OxmlElement cuc bo, tranh loi khi python-docx
# chua cai (None object khong the goi duoc, bi nuot silent trong except).
